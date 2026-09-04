# -*- coding: utf-8 -*-
"""按黑马体例生成：克隆官方源码 + 在 exam_demo 上加完整代码。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path("/workspace/springboot/exam_demo")
OUT_MD = Path("/workspace/3-国家电网数据平台开发步骤.md")
OUT_DOCX = Path("/workspace/3-国家电网数据平台开发步骤.docx")
OUT_MD2 = ROOT / "docs" / "3-国家电网数据平台开发步骤.md"
OUT_DOCX2 = ROOT / "docs" / "3-国家电网数据平台开发步骤.docx"
OUT_TXT = Path("/workspace/项目说明.txt")


def read(rel):
    p = ROOT / rel
    text = p.read_text(encoding="utf-8")
    if text.startswith("\ufeff"):
        text = text[1:]
    return text.rstrip() + "\n"


md = []


def m(s=""):
    md.append(s)


def step_title(n, title):
    m("## %s. %s" % (n, title))
    m()


def h3(title):
    m("### %s" % title)
    m()


def p(s):
    m(s)
    m()


def note(*lines):
    m("**注意：**")
    m()
    for line in lines:
        m("- %s" % line)
    m()


def explain(*lines):
    m("**代码说明：**")
    m()
    for line in lines:
        m("- %s" % line)
    m()


def code(lang, rel):
    m("**文件路径：** `%s`" % rel)
    m()
    m("```%s" % lang)
    m(read(rel).rstrip())
    m("```")
    m()


m("# 3-国家电网数据平台开发步骤")
m()
p("国家电网数据平台必须在试卷给出的官方源码上完成，不要从零新建 state-grid-app。环境：Windows 10/11 + IntelliJ IDEA。考试工程是 Java 8 / Spring Boot 2.6.4；第三方 Blade 是 JDK 17。统一密码 `root / 127307`。")
p("官方地址：")
m("- 考试工程：`https://gitee.com/iocs/exam_demo.git`")
m("- 第三方：`https://gitee.com/qwqweqwe/blade-springboot.git`")
m()
m("开发步骤一览：")
m()
m("1. 克隆两份官方源码并导入 IDEA")
m("2. 启动第三方 blade-springboot")
m("3. 在 exam_demo 增加依赖与 yml")
m("4. 编写配置、实体、Mapper")
m("5. 编写远程调用客户端")
m("6. 编写同步、Excel/MinIO、Word、定时任务")
m("7. 编写 Controller")
m("8. 启动测试")
m("9. 运维 Docker 与 EMP/DEPT SQL")
m()

step_title("1", "克隆官方源码并导入 IDEA")
p("在 `springboot` 目录打开终端（Windows 用 cmd / IDEA Terminal）：")
m("```bat")
m("cd springboot")
m("git clone https://gitee.com/iocs/exam_demo.git")
m("git clone https://gitee.com/qwqweqwe/blade-springboot.git")
m("```")
m()
p("本仓库已经放好这两份源码，对应目录：`springboot\\exam_demo`、`springboot\\blade-springboot`。父工程 `springboot\\pom.xml` 增加模块 `exam_demo`（Blade 独立打开，不要挂到 Java 8 父工程下）。")
note(
    "考试工程启动类是 `com.zhrj.DemoApplication`，不要再写 StateGridApplication。",
    "第三方启动类是 `org.springblade.Application`，端口 9999，已带 `/blade-auth/token` 和 `/blade-user/user-list`。",
    "exam_demo 自带 `context-path: /exam`，试卷接口实际路径是 `/exam/auth/token`。",
)

step_title("2", "启动第三方 blade-springboot")
p("这是试卷说的「第三方系统提供源码」，接口已经写好，学生负责调用，不要再写 blade-mock。")
m("1. 安装 MySQL、Redis。库名 `blade`，账号密码按 `blade-springboot\\src\\main\\resources\\application-dev.yml`（默认 root/123456，Redis 192.168.232.100 请改成 127.0.0.1）。")
m("2. 用客户端执行 `blade-springboot\\doc\\sql\\blade-saber-mysql.sql`。")
m("3. IDEA 打开 `blade-springboot`，SDK 选 JDK 17，运行 `org.springblade.Application`。")
m("4. 控制台出现端口 9999 后，第三方即就绪。clientId=`saber`，clientSecret=`saber_secret` 已在 SQL 的 `blade_client` 表中。")
m()
p("OAuth 调用与试卷一致：")
m("```text")
m("POST http://127.0.0.1:9999/blade-auth/token")
m("Header: Authorization: Basic c2FiZXI6c2FiZXJfc2VjcmV0")
m("Header: Tenant-Id: 000000")
m("Query: tenantId=000000&username=admin&password=21232f297a57a5a743894a0e4a801fc3&grantType=password")
m("```")
m()
p("台账：")
m("```text")
m("GET http://127.0.0.1:9999/blade-user/user-list?current=1&size=10")
m("Header: Blade-Auth: {accessToken}")
m("Header: Authorization: Basic c2FiZXI6c2FiZXJfc2VjcmV0")
m("```")
m()
p("前面加 Nginx 后，把 exam_demo 的 yml 改成试卷原文 `http://localhost/blade-auth/token`。")

step_title("3", "在 exam_demo 增加依赖与配置")
p("官方 pom 已有 Web、MyBatis-Plus、Druid、MySQL。再增加 EasyExcel、MinIO、H2、Actuator。Lombok 指定 1.18.34，避免本机 JDK 21 编译失败。")
h3("3.1 pom.xml（完整）")
code("xml", "pom.xml")
explain(
    "EasyExcel 负责 xlsx，Word 用它带的 Apache POI。",
    "MinIO 账号按试卷：remote_user / 0123456789.abcdefg，bucket=demo。",
    "H2 仅本机无 MySQL 时使用。",
)
h3("3.2 启动类（官方 DemoApplication，加上定时任务）")
code("java", "src/main/java/com/zhrj/DemoApplication.java")
explain(
    "`@MapperScan(\"com.zhrj.exam.mapper\")` 是官方就有的，业务 Mapper 必须放这个包。",
    "`@EnableScheduling` 才能每天 03:00 跑。",
)
h3("3.3 application.yml")
code("yaml", "src/main/resources/application.yml")
h3("3.4 application-druid.yml（官方 MySQL）")
code("yaml", "src/main/resources/application-druid.yml")
h3("3.5 application-h2.yml（本机无 MySQL）")
code("yaml", "src/main/resources/application-h2.yml")
h3("3.6 表结构")
code("sql", "src/main/resources/schema.sql")
p("交付用 MySQL 脚本：")
code("sql", "sql/schema.sql")
note(
    "有 MySQL 时把 `spring.profiles.active` 改成 `druid`，并先在 ry-vue 执行 sql/schema.sql。",
    "DruidDataSourceConfig 加上了 `@Profile(\"druid\")`，避免 H2  profile 时强行注入 Druid。",
)

step_title("4", "编写配置、实体、Mapper")
h3("4.1 远程账号配置")
code("java", "src/main/java/com/zhrj/exam/config/RemoteAuthProperties.java")
h3("4.2 RestTemplate")
code("java", "src/main/java/com/zhrj/exam/config/RestTemplateConfig.java")
h3("4.3 MinIO")
code("java", "src/main/java/com/zhrj/exam/config/MinioProperties.java")
h3("4.4 统一返回")
code("java", "src/main/java/com/zhrj/exam/common/ApiResult.java")
code("java", "src/main/java/com/zhrj/exam/common/GlobalExceptionHandler.java")
h3("4.5 实体与 DTO")
code("java", "src/main/java/com/zhrj/exam/entity/LocalUserLedger.java")
code("java", "src/main/java/com/zhrj/exam/entity/SyncReport.java")
code("java", "src/main/java/com/zhrj/exam/dto/BladeTokenResponse.java")
code("java", "src/main/java/com/zhrj/exam/dto/BladeUserPageResponse.java")
code("java", "src/main/java/com/zhrj/exam/dto/UserLedgerExcel.java")
h3("4.6 Mapper")
code("java", "src/main/java/com/zhrj/exam/mapper/LocalUserLedgerMapper.java")
code("java", "src/main/java/com/zhrj/exam/mapper/SyncReportMapper.java")

step_title("5", "编写远程调用客户端")
h3("5.1 BladeAuthClient")
code("java", "src/main/java/com/zhrj/exam/client/BladeAuthClient.java")
explain(
    "官方参数名是 grantType（同时带 grant_type 兼容）。",
    "Basic 就是 saber:saber_secret。",
    "密码必须是 admin 的 MD5：21232f297a57a5a743894a0e4a801fc3。",
)
h3("5.2 BladeUserClient")
code("java", "src/main/java/com/zhrj/exam/client/BladeUserClient.java")
explain("`Blade-Auth` 填 accessToken 原文，不要加 Bearer。")

step_title("6", "编写同步、Excel/MinIO、Word、定时任务")
h3("6.1 UserLedgerSyncService")
code("java", "src/main/java/com/zhrj/exam/service/UserLedgerSyncService.java")
explain(
    "分页拉完远程用户。本地没有该 id → 新增。",
    "远程 isDeleted 从 0 变 1，或远程列表不再包含该用户 → 记删除。",
    "Excel 先写 data/excel，再尝试上传 MinIO；MinIO 没开只打 warn。",
)
h3("6.2 Word")
code("java", "src/main/java/com/zhrj/exam/service/ReportExportService.java")
m("```text")
m("xxxx年 xx月 xx日 台账统计报告")
m("新增账户：_xx__个")
m("已删除账户：_xx__个")
m("国家电网数据平台")
m("```")
m()
h3("6.3 定时任务")
code("java", "src/main/java/com/zhrj/exam/job/UserLedgerSyncJob.java")
explain("`ApplicationRunner` 启动立刻跑一次；Cron `0 0 3 * * ?` 每天 03:00。")

step_title("7", "编写 Controller")
h3("7.1 GET /auth/token")
code("java", "src/main/java/com/zhrj/exam/controller/AuthController.java")
h3("7.2 GET /report/export 与调试接口")
code("java", "src/main/java/com/zhrj/exam/controller/ReportController.java")
h3("7.3 H2 控制台跳转")
code("java", "src/main/java/com/zhrj/exam/controller/H2ConsoleRedirectController.java")

step_title("8", "启动并测试")
p("顺序：Blade（9999）→ exam_demo（8080）。")
m("```powershell")
m("curl http://127.0.0.1:8080/exam/actuator/health")
m("curl http://127.0.0.1:8080/exam/auth/token")
m("curl -Method POST http://127.0.0.1:8080/exam/job/sync")
m("curl http://127.0.0.1:8080/exam/report/stat?date=2026-09-04")
m("curl -L -o $env:USERPROFILE\\Desktop\\台账统计报告.docx http://127.0.0.1:8080/exam/report/export?date=2026-09-04")
m("```")
m()
p("Apifox 导入 `apifox\\state-grid.postman_collection.json`，baseUrl=`http://127.0.0.1:8080/exam`。")
p("H2 控制台：`http://127.0.0.1:8080/exam/h2-console/` ，JDBC 用项目下 `data/exam`，不要用 `jdbc:h2:~/test`。")

step_title("9", "运维 Docker 与 SQL")
p("试卷要求 CentOS7 + Docker。Windows 可用 Docker Desktop 等价完成，root 密码 127307。")
code("yaml", "docker/docker-compose.yml")
p("EMP / DEPT 题：")
code("sql", "sql/ops-emp-dept.sql")

m("## 附录：端口")
m()
m("| 端口 | 工程 | 用途 |")
m("|---|---|---|")
m("| 9999 | blade-springboot | 官方第三方 OAuth、用户台账 |")
m("| 8080 | exam_demo | 试卷接口（前缀 /exam） |")
m("| 9000/9001 | MinIO | remote_user / 0123456789.abcdefg |")
m("| 3306/3307 | MySQL 主从 | 运维 |")
m("| 6379/26379 | Redis / 哨兵 | 运维 |")
m("| 80 | Nginx | 反代成 localhost/blade-auth |")
m("| 9090/3000 | Prometheus / Grafana | 监控 |")
m()

text = "\n".join(md)
OUT_MD.write_text(text, encoding="utf-8")
OUT_MD2.parent.mkdir(parents=True, exist_ok=True)
OUT_MD2.write_text(text, encoding="utf-8")
OUT_TXT.write_text(read("项目说明.txt"), encoding="utf-8")

doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)


def font(run, cn="宋体", en="Calibri", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), cn)


in_code = False
buf = []


def flush_code():
    global buf
    if not buf:
        return
    lines = buf
    buf = []
    chunk = []
    for line in lines:
        chunk.append(line if line != "" else " ")
        if len(chunk) >= 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            run = p.add_run("\n".join(chunk))
            font(run, "Consolas", "Consolas", 8.5)
            pPr = p._p.get_or_add_pPr()
            pPr.append(parse_xml(
                '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="F4F4F4"/>'
            ))
            chunk = []
    if chunk:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run("\n".join(chunk))
        font(run, "Consolas", "Consolas", 8.5)
        pPr = p._p.get_or_add_pPr()
        pPr.append(parse_xml(
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="F4F4F4"/>'
        ))


for raw in md:
    line = raw
    if line.startswith("```") and not in_code:
        in_code = True
        buf = []
        continue
    if line.startswith("```") and in_code:
        in_code = False
        flush_code()
        continue
    if in_code:
        buf.append(line)
        continue
    if not line.strip():
        continue
    if line.startswith("# "):
        hp = doc.add_heading(line[2:].strip(), 0)
        for r in hp.runs:
            font(r, "黑体", size=22, bold=True, color=(0x1F, 0x4E, 0x79))
        continue
    if line.startswith("## "):
        hp = doc.add_heading(line[3:].strip(), 1)
        for r in hp.runs:
            font(r, "黑体", size=16, bold=True, color=(0x1F, 0x4E, 0x79))
        continue
    if line.startswith("### "):
        hp = doc.add_heading(line[4:].strip(), 2)
        for r in hp.runs:
            font(r, "黑体", size=13, bold=True, color=(0x2E, 0x6B, 0x9A))
        continue
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    text_line = line
    bold = text_line.startswith("**") and text_line.endswith("**") and text_line.count("**") == 2
    run = para.add_run(text_line.replace("**", ""))
    font(run, size=12, bold=bold)

doc.save(OUT_DOCX)
doc.save(OUT_DOCX2)
print("md", OUT_MD, OUT_MD.stat().st_size)
print("docx", OUT_DOCX, OUT_DOCX.stat().st_size)
