# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

out = Path("/workspace/springboot/State-Grid/docs/国家电网数据平台-Windows从零完成教程.docx")
out.parent.mkdir(parents=True, exist_ok=True)

doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.2)
section.right_margin = Cm(2.2)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)


def set_run_font(run, cn="宋体", en="Calibri", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), cn)


def add_heading_cn(text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        size = 18 if level == 1 else 15 if level == 2 else 13
        set_run_font(run, "黑体", "Calibri", size, True, (0x1F, 0x4E, 0x79))
    return p


def add_p(text, size=12, bold=False, first_line=True, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first_line:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run("注意：" + text)
    set_run_font(run, size=11, bold=True, color=(0xB3, 0x3B, 0x1F))
    return p


def add_code(text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(8)
    pf.left_indent = Cm(0.4)
    run = p.add_run(text)
    set_run_font(run, cn="Consolas", en="Consolas", size=10)
    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=11, bold=True, color=(0xFF, 0xFF, 0xFF))
        shd = parse_xml(
            '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="1F4E79"/>'
        )
        cell._tc.get_or_add_tcPr().append(shd)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10.5)
    doc.add_paragraph()
    return table


for _ in range(4):
    doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("国家电网数据平台")
set_run_font(r, "黑体", size=28, bold=True, color=(0x1F, 0x4E, 0x79))

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("Windows 环境从头到尾完成教程")
set_run_font(r, "黑体", size=20, bold=True, color=(0x2E, 0x6B, 0x9A))

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("企业开发项目  +  企业运维项目")
set_run_font(r, size=14, bold=True)

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("对应代码：springboot \\ State-Grid")
set_run_font(r, size=12, color=(0x55, 0x55, 0x55))

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("统一密码 root / 127307")
set_run_font(r, size=12, color=(0x55, 0x55, 0x55))

doc.add_page_break()

add_heading_cn("写在前面：这份教程带你做什么", 1)
add_p("试卷有两份 Word：企业开发项目、企业运维项目。本教程按 Windows 电脑（Windows 10/11 + IntelliJ IDEA）从零做到能交差。")
add_p("开发部分：在 IDEA 里跑两个 Spring Boot，完成拿 Token、同步用户台账、生成 Excel、导出 Word 日报。代码已经写在 springboot\\State-Grid 里，你主要是部署、跑通、截图。")
add_p("运维部分：在 Windows 上用 Docker Desktop 完成 MySQL 主从、Redis 哨兵、Nginx、Prometheus+Grafana；再用客户端做 EMP/DEPT 的 SQL 题。如果老师强制要求 CentOS 7 虚拟机，见第六章。")
add_tip("开发接口必须打 8080（state-grid-app），不要打 18080。18080 只是模拟第三方 Blade。H2 控制台不要用默认的 jdbc:h2:~/test。")

add_heading_cn("一、项目到底要交什么", 1)
add_heading_cn("1. 企业开发项目", 2)
add_p("依据 yml 里配置的租户、账号、密码、clientId/clientSecret，调用远程 OAuth 拿 Token，并提供查询接口返回 Token。")
add_p("用 Token 里的 accessToken 调用用户信息台账分页接口。定时任务启动时跑一次、之后每天 03:00 再跑，增量同步到本地表，同时生成 Excel。")
add_p("提供日报导出：GET /report/export?date={date}，返回 Word，格式为「xxxx年 xx月 xx日 台账统计报告 / 新增账户 / 已删除账户」。")
add_table(
    ["交付物", "本仓库对应位置"],
    [
        ["源码 zip", "把 springboot\\State-Grid 打成 zip（不要含 target、data）"],
        ["数据库 sql（只要结构）", "sql\\schema.sql"],
    ],
)

add_heading_cn("2. 企业运维项目", 2)
add_p("安装 CentOS7（或用 Windows 上的 Docker 等价完成）、Docker、Docker Compose；MySQL 主从热备份；Redis 哨兵；Nginx 反向代理负载均衡；Prometheus+Grafana 监控 Spring Boot；按题目建 EMP/DEPT 并写 SQL。")
add_table(
    ["交付物", "怎么交"],
    [
        ["MySQL 主从安装文档", "本教程第四章可直接当文档"],
        ["Redis 哨兵安装文档", "本教程第四章"],
        ["Nginx 安装文档", "本教程第四章"],
        ["Prometheus+Grafana 安装文档", "本教程第四章"],
        ["数据库 SQL 脚本", "本教程第五章，保存为 ops-emp-dept.sql"],
    ],
)

add_heading_cn("二、Windows 环境准备", 1)
add_heading_cn("1. 必装软件", 2)
add_table(
    ["软件", "建议版本", "用途"],
    [
        ["JDK", "17（Temurin 或 Oracle）", "编译运行 Spring Boot"],
        ["Maven", "3.8 / 3.9", "打包"],
        ["IntelliJ IDEA", "2023 及以上", "打开工程、启动两个应用"],
        ["Apifox", "最新版", "导入并测接口"],
        ["Git（可选）", "最新版", "拉代码"],
        ["Docker Desktop（运维用）", "最新稳定版", "MySQL/Redis/Nginx/监控"],
    ],
)
add_p("JDK 安装完后，系统环境变量要有 JAVA_HOME，例如 C:\\Program Files\\Eclipse Adoptium\\jdk-17，Path 里加上 %JAVA_HOME%\\bin。命令提示符执行：")
add_code("java -version\nmvn -version")
add_p("应能看到 java version \"17...\"。如果提示 mvn 不是内部或外部命令，把 Maven 的 bin 目录也加到 Path。")

add_heading_cn("2. IDEA 打开工程", 2)
add_p("用 IDEA 打开已有工程根目录，确保 springboot\\pom.xml 的 modules 里包含 State-Grid。")
add_p("右侧 Maven 面板点 Reload All Maven Projects。能看到：")
add_code("springboot\n  └─ State-Grid\n       ├─ blade-mock\n       └─ state-grid-app")
add_p("如果 State-Grid 是红的，先确认父 pom 已写入 <module>State-Grid</module>，再 Reload。")

add_heading_cn("3. 两个模块分别干什么", 2)
add_table(
    ["模块", "启动类", "端口", "角色"],
    [
        ["blade-mock", "com.blademock.BladeMockApplication", "18080", "假装第三方 Blade：发 Token、给用户台账"],
        ["state-grid-app", "com.stategrid.StateGridApplication", "8080", "你要交的开发程序：拿 Token、同步、导出 Word"],
    ],
)
add_p("yml 里远程地址默认指向本机 mock，所以不必真去搭 Blade 官方工程。账号如下（已写进 application.yml，不用改）：")
add_table(
    ["配置项", "值"],
    [
        ["clientId", "saber"],
        ["clientSecret", "saber_secret"],
        ["tenantId", "000000"],
        ["username", "admin"],
        ["password", "21232f297a57a5a743894a0e4a801fc3（admin 的 MD5）"],
    ],
)

add_heading_cn("三、开发项目：部署与跑通", 1)
add_heading_cn("1. 业务流程图（演示前先看懂）", 2)
add_code(
    "yml 账号\n   |\n   v\nGET http://127.0.0.1:8080/auth/token     <-- 试卷接口，浏览器/Apifox 打这个\n   |\n   v\nPOST http://127.0.0.1:18080/blade-auth/token   <-- app 内部去调 mock\n   |  Header: Authorization Basic(saber:saber_secret)\n   v\nGET  http://127.0.0.1:18080/blade-user/user-list?current=1&size=10\n   |  Header: Blade-Auth = accessToken\n   v\n增量写入 H2 表 LOCAL_USER_LEDGER\n统计写入 SYNC_REPORT\n生成 Excel 到 state-grid-app\\data\\excel\\\n   v\nGET http://127.0.0.1:8080/report/export?date=yyyy-MM-dd  <-- 下载 Word"
)
add_p("定时：程序启动立刻同步一次；之后每天凌晨 03:00 再同步。演示时用 POST /job/sync 立刻再跑一遍即可。")

add_heading_cn("2. 编译", 2)
add_p("在 IDEA 底部 Terminal（或 Windows 的 cmd / PowerShell）进入目录打包：")
add_code("cd springboot\\State-Grid\nmvn -DskipTests package")
add_p("成功后有两个 jar：")
add_code("blade-mock\\target\\blade-mock-1.0-SNAPSHOT.jar\nstate-grid-app\\target\\state-grid-app-1.0-SNAPSHOT.jar")
add_tip("本机若是 JDK 21，Lombok 已在 pom 里指定 1.18.34，一般可直接编过。仍报错时把 IDEA 的 Project SDK 改成 17 再编。")

add_heading_cn("3. 启动（顺序不能反）", 2)
add_p("方式 A：IDEA 图形界面（推荐）")
add_p("1）打开 BladeMockApplication，点绿色三角 Run。控制台出现 Tomcat started on port(s): 18080 再进行下一步。")
add_p("2）打开 StateGridApplication，再点 Run。出现 Tomcat started on port(s): 8080。同时会打印「应用启动，立即执行一次用户台账同步」。")
add_p("方式 B：两个命令行窗口")
add_code("cd springboot\\State-Grid\\blade-mock\njava -jar target\\blade-mock-1.0-SNAPSHOT.jar")
add_code("cd springboot\\State-Grid\\state-grid-app\njava -jar target\\state-grid-app-1.0-SNAPSHOT.jar")
add_tip("第二个命令的当前目录必须是 state-grid-app，否则 H2 文件路径会错。Windows 不要关这两个窗口。")

add_heading_cn("4. 用浏览器 / PowerShell 验收开发接口", 2)
add_p("PowerShell 里执行（把日期改成当天）：")
add_code(
    "curl http://127.0.0.1:8080/actuator/health\n"
    "curl http://127.0.0.1:8080/auth/token\n"
    "curl -Method POST http://127.0.0.1:8080/job/sync\n"
    "curl http://127.0.0.1:8080/report/stat?date=2026-09-04\n"
    "curl -L -o $env:USERPROFILE\\Desktop\\台账统计报告.docx http://127.0.0.1:8080/report/export?date=2026-09-04"
)
add_table(
    ["接口", "成功长什么样"],
    [
        ["/actuator/health", "status 为 UP，database 为 H2"],
        ["/auth/token", "success true，data.accessToken 有一串字符"],
        ["/job/sync", "totalCount 至少为 1"],
        ["/report/export", "桌面上出现可打开的 docx"],
    ],
)
add_p("Word 打开后应类似：")
add_code("2026年 9月 4日 台账统计报告\n新增账户：_10__个\n已删除账户：_1__个\n国家电网数据平台")

add_heading_cn("5. Apifox 导入与测试", 2)
add_p("1）打开 Apifox → 导入 → Postman。")
add_p("2）选择 springboot\\State-Grid\\apifox\\state-grid.postman_collection.json。")
add_p("3）再导入同目录 state-grid.postman_environment.json。")
add_p("4）环境变量 baseUrl 必须是 http://127.0.0.1:8080 ，不要末尾斜杠，不要写成 18080。")
add_p("5）按文件夹顺序跑：01 获取 Token → 02 台账同步 → 03 导出 Word。05 才是打 18080 的第三方模拟。")
add_tip("若出现 404 且 path 为 //auth/token，就是 baseUrl 多了斜杠或端口用成了 18080。改环境变量后重发即可，不是代码坏了。")

add_heading_cn("6. 用 H2 控制台看库", 2)
add_p("浏览器打开（末尾斜杠不能省）：")
add_code("http://127.0.0.1:8080/h2-console/")
add_p("登录不要用页面默认的 jdbc:h2:~/test（会报 Database not found）。Windows 路径用正斜杠：")
add_code(
    "JDBC URL:\n"
    "jdbc:h2:file:C:/你的路径/springboot/State-Grid/state-grid-app/data/state_grid;MODE=MySQL;AUTO_SERVER=TRUE\n\n"
    "User Name: sa\nPassword: （空）"
)
add_p("例如工程在 D:\\IdeaProjects\\javafirst 下：")
add_code("jdbc:h2:file:D:/IdeaProjects/javafirst/springboot/State-Grid/state-grid-app/data/state_grid;MODE=MySQL;AUTO_SERVER=TRUE")
add_p("Connect 之后应看到两张表：LOCAL_USER_LEDGER（用户台账）、SYNC_REPORT（当天新增/删除统计）。")

add_heading_cn("7. 演示增量同步（答辩加分）", 2)
add_p("在 PowerShell：")
add_code(
    "curl.exe -X POST http://127.0.0.1:18080/blade-user/mock/add -d \"account=newuser01&name=新员工\"\n"
    "curl.exe -X POST http://127.0.0.1:18080/blade-user/mock/delete/10002\n"
    "curl.exe -X POST http://127.0.0.1:8080/job/sync"
)
add_p("第二次同步的 newCount 会加 1，deletedCount 会加 1。再导出一次 Word，两个数字会对上。Excel 在：")
add_code("springboot\\State-Grid\\state-grid-app\\data\\excel\\user-ledger-当天日期.xlsx")

add_heading_cn("8. 开发部分常见问题", 2)
add_table(
    ["现象", "原因", "处理"],
    [
        ["app 启动提示连不上 18080", "mock 没先启动", "先起 blade-mock 再起 app，或再点一次 /job/sync"],
        ["8080 端口占用", "上次 Java 没关", "任务管理器结束 java.exe"],
        ["H2 404", "打开了 /h2-console 没斜杠", "改用 /h2-console/"],
        ["H2 Database not found", "用了默认 ~/test", "换成项目 data\\state_grid 的绝对路径"],
        ["Apifox 404", "打到 18080 或双斜杠", "baseUrl=http://127.0.0.1:8080"],
    ],
)

add_heading_cn("四、运维项目：Windows + Docker Desktop", 1)
add_p("考试原文写的是克隆 CentOS 7。在 Windows 上更稳的做法是安装 Docker Desktop，用容器完成同等拓扑，截图交文档。若机房强制虚拟机，见第六章。")

add_heading_cn("1. 安装 Docker Desktop", 2)
add_p("1）官网下载 Docker Desktop for Windows，安装时勾选 Use WSL 2。")
add_p("2）重启电脑，启动 Docker Desktop，托盘图标不再转圈。")
add_p("3）PowerShell 执行 docker version 和 docker compose version，都能输出版本号。")
add_p("资源建议：Settings → Resources，内存至少 4GB（8GB 更稳）。")

add_heading_cn("2. 启动中间件", 2)
add_p("若仓库有 springboot\\State-Grid\\docker\\docker-compose.yml，直接：")
add_code("cd springboot\\State-Grid\\docker\ndocker compose up -d\ndocker compose ps")
add_p("需要起来的服务与端口：")
add_table(
    ["服务", "端口", "账号"],
    [
        ["MySQL 主", "3306", "root / 127307"],
        ["MySQL 从", "3307", "root / 127307"],
        ["Redis 主", "6379", "密码 127307"],
        ["Redis 从 / 哨兵", "6380、6381 / 26379-26381", "哨兵名 mymaster"],
        ["Nginx", "80", "反代 8080 与 18080"],
        ["MinIO", "9000 / 9001", "remote_user / 0123456789.abcdefg，bucket=demo"],
        ["Prometheus", "9090", "抓取 /actuator/prometheus"],
        ["Grafana", "3000", "admin / 127307"],
    ],
)
add_p("全部 Up 后再做检查。第一次拉镜像需要联网，等几分钟。")

add_heading_cn("3. MySQL 主从验证", 2)
add_code("docker exec -it sg-mysql-slave mysql -uroot -p127307 -e \"SHOW REPLICA STATUS\\G\"")
add_p("看 Replica_IO_Running、Replica_SQL_Running 都是 Yes。再在主库建个测试库，从库能查到即同步成功。也可用 Navicat 连 127.0.0.1:3306（主）和 3307（从）。")

add_heading_cn("4. Redis 哨兵验证", 2)
add_code("docker exec -it sg-redis-sentinel1 redis-cli -p 26379 SENTINEL masters")
add_p("能看到 mymaster。故障转移演示：docker stop 主容器，等约 5 秒再查 get-master-addr-by-name，会切到从库。")

add_heading_cn("5. Nginx 反向代理", 2)
add_p("两个 Spring Boot 保持在 IDEA 里跑。Nginx 把 80 端口转到 8080，并把 /blade-auth/、/blade-user/ 转到 18080。浏览器访问：")
add_code("http://127.0.0.1/auth/token")
add_p("若要演示负载均衡，IDEA 再起一份 state-grid-app，Program arguments 填 --spring.profiles.active=replica（端口 8081）。")
add_tip("Windows 上 80 端口有时被 IIS 占用。占用时把映射改成 8088:80，访问 http://127.0.0.1:8088 。")

add_heading_cn("6. Prometheus + Grafana", 2)
add_p("应用已暴露 http://127.0.0.1:8080/actuator/prometheus 。")
add_p("Prometheus：http://127.0.0.1:9090 → Status → Targets，job=state-grid-app 为 UP。")
add_p("Grafana：http://127.0.0.1:3000 ，账号 admin，密码 127307。数据源选 Prometheus。可导入看板 ID 4701（JVM）。截图：应用 UP、堆内存曲线。")

add_heading_cn("7. MinIO", 2)
add_p("控制台 http://127.0.0.1:9001 ，Access Key=remote_user，Secret Key=0123456789.abcdefg，Bucket=demo。当前精简开发版 Excel 写在本地 data\\excel，答辩时说明即可。")

add_heading_cn("五、运维 SQL 题（直接可交）", 1)
add_p("用客户端连 MySQL（Docker 主库 3306 或本机 MySQL），新建库 ops_demo，执行下面全文，保存为 ops-emp-dept.sql 一并提交。")
add_code(
    "CREATE DATABASE IF NOT EXISTS ops_demo DEFAULT CHARACTER SET utf8mb4;\n"
    "USE ops_demo;\n\n"
    "DROP TABLE IF EXISTS EMP;\nDROP TABLE IF EXISTS DEPT;\n\n"
    "CREATE TABLE DEPT (\n  ID BIGINT PRIMARY KEY,\n  NAME VARCHAR(20) NOT NULL,\n  LOCATION VARCHAR(100)\n);\n\n"
    "CREATE TABLE EMP (\n  ID BIGINT PRIMARY KEY,\n  NAME VARCHAR(100) NOT NULL,\n  SALARY INT,\n  DEPTID INT,\n"
    "  CONSTRAINT fk_emp_dept FOREIGN KEY (DEPTID) REFERENCES DEPT(ID)\n);\n\n"
    "INSERT INTO DEPT VALUES (1,'研发部','一楼'),(2,'人事部','二楼');\n"
    "INSERT INTO EMP VALUES (1001,'张三',5000,1),(1002,'李四',4000,1),(1003,'王五',1000,2),(1004,'赵六',3000,2);\n\n"
    "-- 1 张三薪水提升 10%\nUPDATE EMP SET SALARY = SALARY * 1.1 WHERE NAME = '张三';\n\n"
    "-- 2 姓王的员工按薪水降序\nSELECT * FROM EMP WHERE NAME LIKE '王%' ORDER BY SALARY DESC;\n\n"
    "-- 3 第 2 到 4 条记录\nSELECT * FROM EMP LIMIT 1, 3;\n\n"
    "-- 4 左外连接员工和部门\nSELECT e.*, d.NAME AS DEPT_NAME, d.LOCATION\nFROM EMP e LEFT JOIN DEPT d ON e.DEPTID = d.ID;\n\n"
    "-- 5 各部门人数\nSELECT d.NAME, COUNT(e.ID) AS CNT\nFROM DEPT d LEFT JOIN EMP e ON d.ID = e.DEPTID\nGROUP BY d.ID, d.NAME;\n\n"
    "-- 6 平均薪水大于 3000 的部门\nSELECT d.NAME, AVG(e.SALARY) AS AVG_SALARY\nFROM EMP e JOIN DEPT d ON e.DEPTID = d.ID\nGROUP BY d.ID, d.NAME\nHAVING AVG(e.SALARY) > 3000;"
)

add_heading_cn("六、老师点名要 CentOS 7 虚拟机时（可选）", 1)
add_p("Windows 推荐 VMware Workstation。CentOS 7 选 Minimal 或 DVD ISO。root 密码设为 127307。")
add_p("本机内存建议 16GB 以上。开 2 台即可：")
add_table(
    ["虚拟机", "CPU", "内存", "磁盘", "角色"],
    [
        ["vm1", "2 核", "4GB", "40GB", "Docker、MySQL 主、Nginx、监控、跑应用"],
        ["vm2", "2 核", "4GB", "30GB", "MySQL 从、Redis 哨兵整套"],
    ],
)
add_p("安装完进入系统：systemctl stop firewalld；yum 安装 docker 与 docker-compose；把项目拷进去再 compose up。两台用桥接或仅主机网络，互相能 ping。")
add_tip("笔记本内存小于 8GB 时不要强开两台虚拟机，用第四章 Docker Desktop 即可，答辩说明「Windows 上用容器等价部署」。")

add_heading_cn("七、答辩演示建议顺序（约 15 分钟）", 1)
add_p("1. IDEA 展示两个启动类，先 mock 后 app。")
add_p("2. 浏览器打开 /auth/token，指出 accessToken。")
add_p("3. Apifox 跑 /job/sync，H2 打开 LOCAL_USER_LEDGER。")
add_p("4. 下载 Word 日报，对照新增/删除数字。")
add_p("5. 若运维已起：Docker Desktop 窗口、从库 SHOW REPLICA STATUS、Grafana 看板、Nginx 访问 80 端口。")
add_p("6. 客户端跑第五章 SQL，展示查询结果。")

add_heading_cn("八、打包提交", 1)
add_p("开发：选中 springboot\\State-Grid，排除 target、data、.idea，右键压缩为 State-Grid-src.zip，附上 sql\\schema.sql。")
add_p("运维：本 Word 可作四份安装文档的合并版；SQL 另存 ops-emp-dept.sql。按老师要求拆成四个 txt 也可以，内容从第四章对应小节复制。")
add_p("源码不要提交 H2 的 data 目录和 Excel 运行产物。")

add_heading_cn("附录：端口一览", 1)
add_table(
    ["端口", "谁在听", "你怎么用"],
    [
        ["18080", "blade-mock", "只测 /blade-auth/token、/blade-user/user-list"],
        ["8080", "state-grid-app", "试卷接口、H2、actuator"],
        ["8081", "app 第二实例", "可选，负载均衡"],
        ["3306 / 3307", "MySQL 主/从", "运维"],
        ["6379 等", "Redis / 哨兵", "运维"],
        ["80", "Nginx", "运维反代"],
        ["9000 / 9001", "MinIO", "对象存储控制台"],
        ["9090 / 3000", "Prometheus / Grafana", "监控"],
    ],
)
add_p("做到这里：开发两个进程能跑、四个接口能通、H2 有数据、Word 能下载；运维容器能起、SQL 能跑，即可按试卷要求交付。")

doc.save(out)
print("wrote", out, "bytes", out.stat().st_size)
