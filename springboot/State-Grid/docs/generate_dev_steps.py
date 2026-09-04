# -*- coding: utf-8 -*-
"""按《3-MyBatis开发步骤》体例生成：步骤 + 完整代码 + 代码说明。"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path("/workspace/springboot/State-Grid")
OUT_MD = Path("/workspace/3-国家电网数据平台开发步骤.md")
OUT_DOCX = Path("/workspace/3-国家电网数据平台开发步骤.docx")
OUT_MD2 = ROOT / "docs" / "3-国家电网数据平台开发步骤.md"
OUT_DOCX2 = ROOT / "docs" / "3-国家电网数据平台开发步骤.docx"


def read(rel: str) -> str:
    p = ROOT / rel
    if not p.exists():
        p = Path("/workspace") / rel
    text = p.read_text(encoding="utf-8")
    if text.startswith("\ufeff"):
        text = text[1:]
    return text.rstrip() + "\n"


# ---------- markdown ----------
md = []


def m(s=""):
    md.append(s)


def step_title(n, title):
    m(f"\n## {n}. {title}\n")


def h3(title):
    m(f"\n### {title}\n")


def p(s):
    m(s)
    m()


def note(*lines):
    m("**注意：**")
    m()
    for line in lines:
        m(f"- {line}")
    m()


def explain(*lines):
    m("**代码说明：**")
    m()
    for line in lines:
        m(f"- {line}")
    m()


def code(lang, rel):
    m(f"**文件路径：** `{rel}`")
    m()
    m(f"```{lang}")
    m(read(rel).rstrip())
    m("```")
    m()


m("# 3-国家电网数据平台开发步骤")
m()
p("国家电网数据平台（企业开发 + 企业运维）的开发步骤如下。环境为 **Windows 10/11 + IntelliJ IDEA + JDK 17 + Maven 3.8+**。请严格按照步骤顺序操作：先做第三方模拟工程 `blade-mock`，再做业务工程 `state-grid-app`，最后启动、测试。")
p("对应代码目录：`springboot\\State-Grid`。统一密码：`root / 127307`。")
m("开发步骤一览：")
m()
m("1. 创建 Maven 多模块工程")
m("2. 编写第三方模拟工程 blade-mock（OAuth + 用户台账）")
m("3. 编写业务工程 state-grid-app 的配置与表结构")
m("4. 编写实体类、DTO、Mapper")
m("5. 编写远程调用客户端")
m("6. 编写同步服务、定时任务、Excel/Word 导出")
m("7. 编写 Controller 接口")
m("8. 启动项目并测试")
m("9. 运维 SQL（EMP / DEPT）")
m()

# ===== 1 =====
step_title("1", "创建 Maven 多模块工程")
p("打开 IDEA，打开已有的 `springboot` 父工程（或 File → Open 选择工程根目录）。在父工程 `springboot\\pom.xml` 的 `<modules>` 中增加 `State-Grid`。")
m("父工程片段：")
m()
m("```xml")
m("        <module>satoken</module>")
m("        <module>State-Grid</module>")
m("```")
m()
p("在 `springboot` 目录下新建文件夹 `State-Grid`，再在其中创建聚合模块 pom。")
h3("1.1 State-Grid 聚合 pom")
code("xml", "pom.xml")
explain(
    "`packaging` 为 `pom`，表示这是聚合工程，本身不产出 jar。",
    "两个子模块：`state-grid-app`（试卷要求的开发程序）、`blade-mock`（模拟第三方 Blade 接口）。",
)
p("目录结构最终如下：")
m("```text")
m("springboot")
m(" └─ State-Grid")
m("     ├─ pom.xml")
m("     ├─ blade-mock")
m("     │   ├─ pom.xml")
m("     │   └─ src/main/java/com/blademock/...")
m("     └─ state-grid-app")
m("         ├─ pom.xml")
m("         └─ src/main/java/com/stategrid/...")
m("```")
m()
note(
    "在 IDEA 右侧 Maven 面板点击 Reload All Maven Projects。",
    "若模块是红色，检查父 pom 是否已加入 `<module>State-Grid</module>`。",
)

# ===== 2 =====
step_title("2", "编写第三方模拟工程 blade-mock")
p("试卷要求调用 `POST /blade-auth/token` 和 `GET /blade-user/user-list`。本工程用 `blade-mock` 在本机 18080 端口模拟这两套接口，这样不必去搭真实 Blade。")
h3("2.1 添加依赖（blade-mock/pom.xml）")
code("xml", "blade-mock/pom.xml")
explain(
    "父工程使用 `spring-boot-starter-parent` 2.6.2，与当前 springboot 目录其它模块一致。",
    "`spring-boot-starter-web` 提供 Tomcat 和 Spring MVC。",
    "Lombok 指定 1.18.34，避免 JDK 21 编译时报 `JCImport qualid`。",
)

h3("2.2 编写启动类")
p("在 `com.blademock` 包下创建 `BladeMockApplication`。")
code("java", "blade-mock/src/main/java/com/blademock/BladeMockApplication.java")

h3("2.3 编写 application.yml")
p("在 `src/main/resources` 下创建 `application.yml`，端口固定 18080。")
code("yaml", "blade-mock/src/main/resources/application.yml")

h3("2.4 编写内存数据仓库 MockDataStore")
p("在 `com.blademock.store` 包下创建。用内存 Map 保存 Token 和用户，启动时预置 10 个用户（其中郑十 isDeleted=1，用于验证删除统计）。")
code("java", "blade-mock/src/main/java/com/blademock/store/MockDataStore.java")
explain(
    "`@PostConstruct init()` 在容器启动后插入演示用户。",
    "`issueToken` 生成 UUID 作为 accessToken，有效期 3600 秒。",
    "`page` 按 current/size 分页，返回结构与 Blade 的 records/total/pages 一致。",
    "`add` / `logicalDelete` 供增量同步演示。",
)

h3("2.5 编写 OAuth 接口 AuthTokenController")
p("对应试卷：`POST http://localhost/blade-auth/token`，Header 为 `Authorization: Basic base64(saber:saber_secret)`，URL 参数含 tenantId、username、password（MD5）、grant_type、type、scope。")
code("java", "blade-mock/src/main/java/com/blademock/controller/AuthTokenController.java")
explain(
    "`saber:saber_secret` 做 Base64 后是 `c2FiZXI6c2FiZXJfc2VjcmV0`。",
    "密码必须是 `21232f297a57a5a743894a0e4a801fc3`（字符串 admin 的 MD5）。",
    "校验通过后返回 `data.accessToken`，业务工程只取这一字段去调台账。",
)

h3("2.6 编写用户台账接口 UserListController")
p("对应试卷：`GET /blade-user/user-list?current=1&size=10`，Header 要带 `Authorization Basic` 和 `Blade-Auth: accessToken`。")
code("java", "blade-mock/src/main/java/com/blademock/controller/UserListController.java")
explain(
    "`isDeleted=1` 表示逻辑删除，同步时计入「已删除账户」。",
    "`/mock/add`、`/mock/delete/{id}` 不是试卷接口，仅用于答辩时演示增量。",
)

# ===== 3 =====
step_title("3", "编写业务工程配置与表结构")
p("`state-grid-app` 才是试卷要交的开发程序，端口 8080。")
h3("3.1 添加依赖（state-grid-app/pom.xml）")
code("xml", "state-grid-app/pom.xml")
explain(
    "`mybatis-plus-boot-starter` 操作本地台账表。",
    "`h2` 作为本机数据库，Windows 上不必先装 MySQL。",
    "`easyexcel` 生成 xlsx；Word 使用 EasyExcel 自带的 Apache POI。",
    "`actuator` + `micrometer-registry-prometheus` 给运维监控用。",
)

h3("3.2 编写启动类")
p("在 `com.stategrid` 包下创建。开启定时任务、扫描 Mapper、绑定远程配置。")
code("java", "state-grid-app/src/main/java/com/stategrid/StateGridApplication.java")
explain(
    "`@EnableScheduling` 才能让 `@Scheduled` 每天 03:00 执行。",
    "`@MapperScan(\"com.stategrid.mapper\")` 相当于给每个 Mapper 加 `@Mapper`。",
)

h3("3.3 编写 application.yml")
p("账户信息必须写在 yml 中，这是试卷明确要求。")
code("yaml", "state-grid-app/src/main/resources/application.yml")
explain(
    "`remote.oauth.*` 就是试卷里的租户、账号、clientId、clientSecret。",
    "默认直连 `127.0.0.1:18080` 的 mock；若前面加了 Nginx，可改成 `http://127.0.0.1/blade-auth/token`。",
    "H2 控制台路径 `/h2-console`，`web-allow-others: true` 允许浏览器访问。",
    "`matching-strategy: ant_path_matcher` 避免 Spring Boot 2.6 下 H2 控制台 404。",
)

h3("3.4 编写表结构 schema.sql")
p("放在 `src/main/resources/schema.sql`，应用启动时自动建表。交付物 `sql/schema.sql` 内容相同（只要结构）。")
code("sql", "state-grid-app/src/main/resources/schema.sql")
explain(
    "`local_user_ledger.id` 使用远程用户 ID，便于增量对比。",
    "`is_deleted`：0 正常，1 逻辑删除。",
    "`sync_report` 按日期唯一，保存当天新增数、删除数，供 Word 导出。",
)

h3("3.5 远程配置属性类")
code("java", "state-grid-app/src/main/java/com/stategrid/config/RemoteAuthProperties.java")
h3("3.6 RestTemplate 配置")
code("java", "state-grid-app/src/main/java/com/stategrid/config/RestTemplateConfig.java")
h3("3.7 统一返回与异常处理")
code("java", "state-grid-app/src/main/java/com/stategrid/common/ApiResult.java")
code("java", "state-grid-app/src/main/java/com/stategrid/common/GlobalExceptionHandler.java")

# ===== 4 =====
step_title("4", "编写实体类、DTO、Mapper")
h3("4.1 本地台账实体 LocalUserLedger")
p("在 `com.stategrid.entity` 包下创建，对应表 `local_user_ledger`。")
code("java", "state-grid-app/src/main/java/com/stategrid/entity/LocalUserLedger.java")
explain(
    "`@TableId(type = IdType.INPUT)` 表示主键由远程 ID 传入，不要本地自增。",
    "字段驼峰对应表字段下划线，由 MyBatis-Plus `map-underscore-to-camel-case` 完成。",
)
h3("4.2 同步日报实体 SyncReport")
code("java", "state-grid-app/src/main/java/com/stategrid/entity/SyncReport.java")
h3("4.3 远程 Token / 用户分页 DTO")
code("java", "state-grid-app/src/main/java/com/stategrid/dto/BladeTokenResponse.java")
code("java", "state-grid-app/src/main/java/com/stategrid/dto/BladeUserPageResponse.java")
h3("4.4 Excel 导出行 UserLedgerExcel")
code("java", "state-grid-app/src/main/java/com/stategrid/dto/UserLedgerExcel.java")
h3("4.5 Mapper 接口")
p("MyBatis-Plus 的 `BaseMapper` 已提供 insert/update/selectById/selectList，不必写 XML。")
code("java", "state-grid-app/src/main/java/com/stategrid/mapper/LocalUserLedgerMapper.java")
code("java", "state-grid-app/src/main/java/com/stategrid/mapper/SyncReportMapper.java")

# ===== 5 =====
step_title("5", "编写远程调用客户端")
h3("5.1 BladeAuthClient：去 mock 拿 Token")
code("java", "state-grid-app/src/main/java/com/stategrid/client/BladeAuthClient.java")
explain(
    "请求方式与试卷一致：POST + Basic + URL 参数。",
    "内存缓存 Token，过期前 60 秒重新拉取。",
)
h3("5.2 BladeUserClient：带 Blade-Auth 拉台账")
code("java", "state-grid-app/src/main/java/com/stategrid/client/BladeUserClient.java")
explain(
    "`Content-Type` 使用试卷要求的 `application/json;charset=UTF-8`。",
    "`Blade-Auth` 填 accessToken 原文，不要加 Bearer。",
)

# ===== 6 =====
step_title("6", "编写同步服务、定时任务、导出")
h3("6.1 UserLedgerSyncService（核心增量逻辑）")
code("java", "state-grid-app/src/main/java/com/stategrid/service/UserLedgerSyncService.java")
explain(
    "分页把远程用户全部拉完。",
    "本地没有该 id → INSERT，新增 +1；若远程已经 isDeleted=1，删除也 +1。",
    "本地有、且 isDeleted 从 0 变为 1 → 删除 +1。",
    "同一天多次同步，日报数字累加本次增量。",
    "Excel 写到运行目录 `data/excel/user-ledger-yyyy-MM-dd.xlsx`。",
)
h3("6.2 Word 导出 ReportExportService")
code("java", "state-grid-app/src/main/java/com/stategrid/service/ReportExportService.java")
p("生成的 Word 格式与试卷一致：")
m("```text")
m("xxxx年 xx月 xx日 台账统计报告")
m("新增账户：_xx__个")
m("已删除账户：_xx__个")
m("国家电网数据平台")
m("```")
m()
h3("6.3 定时任务 UserLedgerSyncJob")
p("试卷要求：程序启动时自动运行一次，之后 Cron 每天 03:00:00。")
code("java", "state-grid-app/src/main/java/com/stategrid/job/UserLedgerSyncJob.java")
explain(
    "`ApplicationRunner` 在 Spring 启动完成后立刻执行。",
    "Cron `0 0 3 * * ?` 是 Spring 六位表达式：秒 分 时 日 月 周。",
    "启动时 mock 若还没起来，只打 warn，等 mock 起来后手动 POST /job/sync。",
)

# ===== 7 =====
step_title("7", "编写 Controller 接口")
h3("7.1 试卷接口 GET /auth/token")
code("java", "state-grid-app/src/main/java/com/stategrid/controller/AuthController.java")
h3("7.2 试卷接口 GET /report/export 以及调试接口")
code("java", "state-grid-app/src/main/java/com/stategrid/controller/ReportController.java")
explain(
    "`GET /report/export?date=yyyy-MM-dd` 返回 docx 附件。",
    "`POST /job/sync` 不是试卷必交接口，用来立刻同步，不必等到凌晨 3 点。",
    "`GET /report/stat` 用 JSON 看当天统计，方便 Apifox 断言。",
)
h3("7.3 H2 控制台斜杠跳转")
code("java", "state-grid-app/src/main/java/com/stategrid/controller/H2ConsoleRedirectController.java")

# ===== 8 =====
step_title("8", "启动项目并测试（Windows）")
h3("8.1 启动顺序")
p("必须先启动 mock，再启动 app。")
m("1. IDEA 打开 `BladeMockApplication`，点击 Run，控制台出现 `Tomcat started on port(s): 18080`。")
m("2. 再打开 `StateGridApplication`，点击 Run，出现 `Tomcat started on port(s): 8080`。")
m("3. 日志中应有「应用启动，立即执行一次用户台账同步」。")
m()
p("或使用两个 cmd 窗口：")
m("```bat")
m("cd springboot\\State-Grid\\blade-mock")
m("java -jar target\\blade-mock-1.0-SNAPSHOT.jar")
m()
m("cd springboot\\State-Grid\\state-grid-app")
m("java -jar target\\state-grid-app-1.0-SNAPSHOT.jar")
m("```")
m()
note("第二个命令的当前目录必须是 state-grid-app，H2 文件才会写到该目录下的 data 文件夹。")

h3("8.2 接口测试（PowerShell）")
m("```powershell")
m("curl http://127.0.0.1:8080/actuator/health")
m("curl http://127.0.0.1:8080/auth/token")
m("curl -Method POST http://127.0.0.1:8080/job/sync")
m("curl http://127.0.0.1:8080/report/stat?date=2026-09-04")
m("curl -L -o $env:USERPROFILE\\Desktop\\台账统计报告.docx http://127.0.0.1:8080/report/export?date=2026-09-04")
m("```")
m()
p("**Apifox：** 导入 `springboot\\State-Grid\\apifox\\state-grid.postman_collection.json`。环境变量 `baseUrl` 必须是 `http://127.0.0.1:8080`，不要斜杠，不要写成 18080。")
note(
    "打到 18080 的 `/auth/token` 会 404，path 变成 `//auth/token` 也是 404。",
    "试卷接口在 8080：`/auth/token`、`/report/export`。",
)

h3("8.3 H2 控制台")
p("浏览器打开（末尾斜杠不能省）：`http://127.0.0.1:8080/h2-console/`")
p("不要用默认 JDBC URL `jdbc:h2:~/test`。改成（注意 Windows 用正斜杠）：")
m("```text")
m("jdbc:h2:file:D:/IdeaProjects/你的工程/springboot/State-Grid/state-grid-app/data/state_grid;MODE=MySQL;AUTO_SERVER=TRUE")
m("User Name: sa")
m("Password: （空）")
m("```")
m()
p("能看到表 `LOCAL_USER_LEDGER`、`SYNC_REPORT` 即为成功。")

h3("8.4 增量演示")
m("```powershell")
m("curl.exe -X POST http://127.0.0.1:18080/blade-user/mock/add -d \"account=newuser01&name=新员工\"")
m("curl.exe -X POST http://127.0.0.1:18080/blade-user/mock/delete/10002")
m("curl.exe -X POST http://127.0.0.1:8080/job/sync")
m("```")
m()
p("第二次同步 newCount、deletedCount 会各加 1。Excel 在 `state-grid-app\\data\\excel\\`。")

# ===== 9 =====
step_title("9", "运维 SQL（EMP / DEPT）")
p("将下面完整脚本保存为 `sql\\ops-emp-dept.sql`，用 Navicat 连接 MySQL 执行。无 MySQL 时，开发阶段可只交该文件。")
code("sql", "sql/ops-emp-dept.sql")
explain(
    "LIMIT 1,3 表示从第 2 条起取 3 条（即第 2～4 条）。",
    "张三薪水提升 10% 后，研发部平均薪水大于 3000，人事部平均 2000 不会出现在第 6 题结果中。",
)

m("\n## 附录：端口一览\n")
m("| 端口 | 工程 | 用途 |")
m("|---|---|---|")
m("| 18080 | blade-mock | 第三方 OAuth、用户台账 |")
m("| 8080 | state-grid-app | 试卷接口、H2、监控 |")
m()
p("做到第 8 步四个接口通、H2 有数据、Word 能下载，开发项目即完成。源码 zip 排除 `target` 和 `data`，附上 `sql\\schema.sql` 提交。")

text = "\n".join(md)
OUT_MD.write_text(text, encoding="utf-8")
OUT_MD2.parent.mkdir(parents=True, exist_ok=True)
OUT_MD2.write_text(text, encoding="utf-8")

# ---------- Word ----------
doc = Document()
sec = doc.sections[0]
sec.page_width = Cm(21)
sec.page_height = Cm(29.7)
sec.left_margin = Cm(2.0)
sec.right_margin = Cm(2.0)
sec.top_margin = Cm(2.0)
sec.bottom_margin = Cm(2.0)


def font(run, cn="宋体", en="Calibri", size=12, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = en
    if color:
        run.font.color.rgb = RGBColor(*color)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), cn)


def shade_para(p):
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(
        '<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:fill="F4F4F4"/>'
    )
    pPr.append(shd)


in_code = False
code_lang = ""
buf = []


def flush_code():
    global buf
    if not buf:
        return
    block = "\n".join(buf)
    buf = []
    # split long blocks into paragraphs of ~80 lines to keep Word happy
    lines = block.split("\n")
    chunk = []
    for line in lines:
        chunk.append(line if line != "" else " ")
        if len(chunk) >= 80:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Cm(0.3)
            run = p.add_run("\n".join(chunk))
            font(run, "Consolas", "Consolas", 8.5)
            shade_para(p)
            chunk = []
    if chunk:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run("\n".join(chunk))
        font(run, "Consolas", "Consolas", 8.5)
        shade_para(p)


for raw in md:
    line = raw
    if line.startswith("```") and not in_code:
        in_code = True
        buf = []
        continue
    if line.startswith("```") and in_code:
        in_code = False
        flush_code()
        continue
    if in_code:
        buf.append(line)
        continue
    if not line.strip():
        continue
    if line.startswith("# "):
        p = doc.add_heading(line[2:].strip(), 0)
        for r in p.runs:
            font(r, "黑体", size=22, bold=True, color=(0x1F, 0x4E, 0x79))
        continue
    if line.startswith("## "):
        p = doc.add_heading(line[3:].strip(), 1)
        for r in p.runs:
            font(r, "黑体", size=16, bold=True, color=(0x1F, 0x4E, 0x79))
        continue
    if line.startswith("### "):
        p = doc.add_heading(line[4:].strip(), 2)
        for r in p.runs:
            font(r, "黑体", size=13, bold=True, color=(0x2E, 0x6B, 0x9A))
        continue
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    text_line = line
    bold = text_line.startswith("**") and text_line.endswith("**") and text_line.count("**") == 2
    if bold:
        run = p.add_run(text_line.replace("**", ""))
        font(run, size=12, bold=True)
    else:
        run = p.add_run(text_line.replace("**", ""))
        font(run, size=12)

doc.save(OUT_DOCX)
doc.save(OUT_DOCX2)
print("md", OUT_MD, OUT_MD.stat().st_size)
print("docx", OUT_DOCX, OUT_DOCX.stat().st_size)
